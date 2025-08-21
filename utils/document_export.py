from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import re

def export_to_txt(report_content, filename="research_report.txt"):
    """Export research report to text file"""
    return report_content, filename

def export_to_docx(report_content, filename="research_report.docx"):
    """Export research report to Word document"""
    doc = Document()
    
    
    title = doc.add_heading('Research Report', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    
    lines = report_content.split('\n')
    
    for line in lines:
        if not line.strip():
            continue
            
        
        if re.match(r'^#+ ', line):
            
            heading_level = min(line.count('#'), 3)  
            heading_text = re.sub(r'^#+ ', '', line)
            doc.add_heading(heading_text, level=heading_level)
        else:
            
            doc.add_paragraph(line)
    
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue(), filename

def export_to_pdf(report_content, filename="research_report.pdf"):
    """Export research report to PDF"""
    
    buffer = BytesIO()
    
    
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
   
    title = Paragraph("Research Report", styles['Title'])
    story.append(title)
    story.append(Spacer(1, 12))
    
    
    lines = report_content.split('\n')
    
    for line in lines:
        if not line.strip():
            story.append(Spacer(1, 6))
            continue
            
       
        if re.match(r'^#+ ', line):
            
            heading_level = min(line.count('#'), 3)  
            heading_text = re.sub(r'^#+ ', '', line)
            story.append(Paragraph(heading_text, styles[f'Heading{heading_level}']))
        else:
           
            story.append(Paragraph(line, styles['BodyText']))
    
    # Build PDF
    doc.build(story)
    
    
    pdf = buffer.getvalue()
    buffer.close()
    
    return pdf, filename